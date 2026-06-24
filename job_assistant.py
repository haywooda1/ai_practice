import anthropic
import os
from datetime import datetime
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

load_dotenv()
client = anthropic.Anthropic()

# ── Token usage tracker ──────────────────────────────────────────
token_log = {"input": 0, "output": 0}

def log_usage(response):
    """Accumulate token usage from any API response."""
    token_log["input"]  += response.usage.input_tokens
    token_log["output"] += response.usage.output_tokens

def print_token_summary():
    """Print token usage and estimated cost at end of run."""
    input_cost  = token_log["input"]  / 1_000_000 * 3.00
    output_cost = token_log["output"] / 1_000_000 * 15.00
    total_cost  = input_cost + output_cost
    print(f"\n{'='*60}")
    print(f"  TOKEN USAGE")
    print(f"{'='*60}")
    print(f"  Input:  {token_log['input']:,} tokens  (${input_cost:.4f})")
    print(f"  Output: {token_log['output']:,} tokens  (${output_cost:.4f})")
    print(f"  Total:  ${total_cost:.4f} this run  ({token_log['input'] // 1000 + 1} letters generated)")
    print(f"  Model:  claude-sonnet-4-5")
# ─────────────────────────────────────────────────────────────────

# Folder where cover letters get saved
OUTPUT_DIR = "cover_letters"
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Your background — pulled directly from your resume
MY_BACKGROUND = """
Name: Adam Haywood
Location: Raleigh-Durham, NC
Contact: unixnrdu@yahoo.com | (919)880-8528 | linkedin.com/in/haywooda
Resume Title: Engineering Leader — Platform Delivery, Reliability & Engineering Operations
Identity: U.S. Navy Veteran | 25+ Years Experience

TARGET ROLES (in priority order):
1. Senior Engineering Manager — Platform, DevOps, SRE, Quality Engineering, or Enablement (PRIMARY TARGET)
2. Director of Engineering — Platform Delivery, Reliability, or Engineering Operations (upside, not required)
3. Director of Quality Engineering / Engineering Enablement (upside, not required)
4. Head of Engineering Operations or Technical Program Delivery

EXECUTIVE SUMMARY:
Engineering leader with 25+ years of experience leading platform delivery, production
reliability, quality engineering, and partner enablement at enterprise scale. Owned the
engineering operations and delivery governance that directly enabled $94.5M in quarterly
revenue. Proven track record building and scaling high-performing engineering teams,
operationalizing AI/ML-enabled automation, and establishing the SDLC governance, CI/CD
practices, and operational excellence standards that ensure production-grade systems
perform under load. Strong credibility in distributed systems, cross-functional
leadership, and engineering change management — with a hands-on approach to both
people and craft.

CORE COMPETENCIES:
- Platform Engineering & Delivery Leadership
- Production Reliability & Site Operations
- Quality Engineering at Enterprise Scale
- SDLC Governance & Release Engineering
- CI/CD Pipelines & Engineering Automation
- AI/ML-Enabled Engineering Modernization
- Capex/Opex Planning & Vendor Governance
- Engineering Talent Strategy & Org Development
- Cross-Functional Leadership: Product, Program, Architecture
- Distributed Systems, APIs & Multi-Platform Integration

PROFESSIONAL EXPERIENCE:

NetApp | Raleigh-Durham, NC | Jan 2005 – May 2026

Senior Manager, Partner & Customer Enablement Solutions | June 2021 – May 2026
- Owned engineering delivery and partner readiness governance that directly enabled
  $94.5M in Q1/Q2 quarterly revenue.
- Led strategy and operationalization of AI/ML-enabled automation including multi-agent
  failure triage — reducing manual diagnostic effort and increasing testing velocity.
- Built and led a globally distributed organization of 10 FTEs and 18 contingent
  engineers across multiple time zones.
- Established delivery governance frameworks, quality standards, and release controls
  that improved execution consistency and maintained platform reliability.
- Owned Capex/Opex planning, SOW negotiation, and vendor oversight.
- Served on cross-functional portfolio councils (FPVR, PPMF) driving investment
  prioritization and risk governance for revenue-critical platform releases.

Manager, Partner & Customer Enablement Solutions | June 2011 – June 2021
- Scaled and led quality engineering capabilities supporting SAN/NAS storage and
  distributed system technologies — sustaining 99.9% interoperability standards.
- Established standardized R&D and quality engineering workflows and release controls.
- Directed validation engineering for distributed storage platforms and API-driven
  integration architectures.
- Partnered with R&D, Product Management, and external technology partners.

Team Lead, Rapid Response Engineering (RRE) — QA & Interoperability | Jan 2005 – May 2011
- Led production incident response, root cause analysis, and platform interoperability
  engineering for customer-critical multi-platform environments.
- Architected SAN test strategies across Solaris, RHEL, AIX, HP-UX, Windows.
- Introduced automation to streamline lab provisioning, CI pipeline workflows, and
  test execution.

MILITARY SERVICE:
United States Navy — Electronics Technician, Honorably Discharged
Specialized in Satellite and Radar equipment, microwave communications, and
navigation systems.

CURRENTLY LEARNING:
Hands-on AI development — Claude API, MCP servers, Python agents, GitHub.
Actively building AI-powered tools and agents.
"""

def get_job_description():
    print("\n📋 Paste the job description below.")
    print("When done, type END on its own line and press Enter:\n")
    lines = []
    while True:
        line = input()
        if line.strip() == "END":
            break
        lines.append(line)
    return "\n".join(lines)


def get_job_title():
    return input("\n📌 Enter a short job title for the filename (e.g. 'ai-eng-manager-google'): ").strip()


def generate_cover_letter(job_description):
    print("\n⏳ Generating your cover letter...")
    message = client.messages.create(
        model="claude-sonnet-4-5",
        max_tokens=1024,
        system="""You are an expert career coach and cover letter writer specializing
in platform delivery, engineering reliability, and engineering operations leadership
roles. Write a tailored, compelling cover letter based on the candidate's background
and the job description provided.

Guidelines:
- Be specific — reference actual requirements from the job posting by name
- Lead with platform delivery, reliability, or engineering operations framing as the
  primary hook — the $94.5M revenue impact and AI/ML automation work are supporting
  proof points, not the headline
- Match the seniority signaled by the job posting — do not over-claim Director/VP
  scope if the role reads as Senior Manager level, and vice versa
- Mention Navy veteran background if the role values leadership or discipline
- Keep it to 3-4 tight paragraphs — hiring managers don't read long letters
- Professional but warm tone — avoid stiff corporate language
- Never use generic filler phrases like 'I am excited to apply' or 'I am a team player'
- End with a confident, specific call to action""",
        messages=[{
            "role": "user",
            "content": f"My background:\n{MY_BACKGROUND}\n\nJob description:\n{job_description}"
        }]
    )
    log_usage(message)  # ← track tokens
    return message.content[0].text


# Save document with a filename that includes the job title and timestamp for easy reference.
#  Include the job description in the saved file for context when reviewing later. .TXT format
#
# def save_cover_letter(cover_letter, job_description, job_title):
#    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
#    safe_title = job_title.replace(" ", "-").lower()
#    filename = f"{OUTPUT_DIR}/{timestamp}_{safe_title}.txt"
#    with open(filename, "w") as f:
#        f.write(f"JOB: {job_title}\n")
#        f.write(f"DATE: {datetime.now().strftime('%B %d, %Y')}\n")
#        f.write(f"{'='*60}\n\n")
#        f.write(f"COVER LETTER:\n\n{cover_letter}\n\n")
#        f.write(f"{'='*60}\n\n")
#        f.write(f"JOB DESCRIPTION:\n\n{job_description}")
#    return filename


 
# ── CHANGED: save_cover_letter now builds a .docx instead of a .txt ──────────
#
# What changed and why:
#
#   OLD: opened a plain text file with open(filename, "w") and wrote
#        everything as a string with f.write().
#
#   NEW: uses python-docx to build a structured Word document with:
#        - A styled header block (name, title, contact info)
#        - A horizontal rule divider
#        - The generated cover letter body as formatted paragraphs
#        - A metadata section at the end with job title and date
#        - The raw job description appended for reference
#
#   The Document object works like a canvas — you add paragraphs
#   one at a time and style each one individually (font, size, color,
#   alignment, bold, spacing). Nothing is written to disk until
#   doc.save(filename) is called at the very end.
#
def save_cover_letter(cover_letter, job_description, job_title):
    # Build the output filename — same naming pattern as before, just .docx now
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_title = job_title.replace(" ", "-").lower()
    filename = f"{OUTPUT_DIR}/{timestamp}_{safe_title}.docx"   # <-- .docx not .txt
 
    # Create a blank Word document
    doc = Document()
 
    # ── Page margins ─────────────────────────────────────────────────────────
    # python-docx sizes use "twips" (1 inch = 914400 EMUs, but for margins
    # we use the Pt() or Inches() helpers). Here we tighten margins slightly
    # so the letter has more breathing room — matching our polished .docx style.
    for section in doc.sections:
        section.top_margin    = Inches(0.85)
        section.bottom_margin = Inches(0.85)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)
 
    # ── Helper: set font on every run in a paragraph ──────────────────────────
    # A "run" in python-docx is a contiguous stretch of text with the same
    # formatting. When you do para.add_run("text"), you get one run back.
    # This helper styles it consistently so we don't repeat ourselves.
    def style_run(run, size=11, bold=False, color=None):
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            # RGBColor takes three ints: red, green, blue (0-255)
            run.font.color.rgb = RGBColor(*color)
 
    # ── NAME (large, dark blue, bold) ────────────────────────────────────────
    name_para = doc.add_paragraph()
    name_para.paragraph_format.space_after = Pt(2)
    name_run = name_para.add_run("ADAM HAYWOOD")
    style_run(name_run, size=20, bold=True, color=(31, 56, 100))   # dark navy
 
    # ── Sub-title ─────────────────────────────────────────────────────────────
    subtitle_para = doc.add_paragraph()
    subtitle_para.paragraph_format.space_after = Pt(2)
    sub_run = subtitle_para.add_run(
        "Engineering Leader — Platform Delivery, Reliability & Engineering Operations"
    )
    style_run(sub_run, size=10, color=(85, 85, 85))
 
    # ── Contact line ──────────────────────────────────────────────────────────
    contact_para = doc.add_paragraph()
    contact_para.paragraph_format.space_after = Pt(2)
    contact_run = contact_para.add_run(
        "unixnrdu@yahoo.com  |  (919) 880-8528  |  Raleigh-Durham, NC  |  U.S. Navy Veteran"
    )
    style_run(contact_run, size=9, color=(85, 85, 85))
 
    # ── LinkedIn ──────────────────────────────────────────────────────────────
    linkedin_para = doc.add_paragraph()
    linkedin_para.paragraph_format.space_after = Pt(6)
    linkedin_run = linkedin_para.add_run("linkedin.com/in/haywooda")
    style_run(linkedin_run, size=9, color=(85, 85, 85))
 
    # ── Horizontal rule (a paragraph with a bottom border) ───────────────────
    # python-docx doesn't have a built-in "insert horizontal rule" method,
    # so we fake it by applying a bottom border to an empty paragraph.
    # The border XML is injected directly into the paragraph's properties.
    rule_para = doc.add_paragraph()
    rule_para.paragraph_format.space_after = Pt(10)
    pPr = rule_para._p.get_or_add_pPr()     # get the paragraph properties element
    pBdr = OxmlElement("w:pBdr")             # create a border container
    bottom = OxmlElement("w:bottom")         # create the bottom border
    bottom.set(qn("w:val"), "single")        # solid line style
    bottom.set(qn("w:sz"), "6")              # thickness (half-points, so 6 = 0.75pt)
    bottom.set(qn("w:space"), "1")           # space between text and border
    bottom.set(qn("w:color"), "1F3864")      # dark navy to match the name color
    pBdr.append(bottom)
    pPr.append(pBdr)
 
    # ── Date ─────────────────────────────────────────────────────────────────
    date_para = doc.add_paragraph()
    date_para.paragraph_format.space_after = Pt(8)
    date_run = date_para.add_run(datetime.now().strftime("%B %d, %Y"))
    style_run(date_run, size=10, color=(85, 85, 85))
 
    # ── Salutation ────────────────────────────────────────────────────────────
    sal_para = doc.add_paragraph()
    sal_para.paragraph_format.space_after = Pt(8)
    sal_run = sal_para.add_run("Dear Hiring Manager,")
    style_run(sal_run, size=11)
 
    # ── Cover letter body ─────────────────────────────────────────────────────
    # The AI returns the letter as a block of text with paragraphs separated
    # by blank lines. We split on double newlines so each paragraph gets its
    # own Word paragraph (proper spacing) rather than one giant text blob.
    body_paragraphs = [p.strip() for p in cover_letter.split("\n\n") if p.strip()]
 
    for para_text in body_paragraphs:
        # Collapse any internal line breaks within a paragraph into spaces
        para_text = " ".join(para_text.splitlines())
        body_para = doc.add_paragraph()
        body_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY   # justified like a real letter
        body_para.paragraph_format.space_after = Pt(8)
        body_run = body_para.add_run(para_text)
        style_run(body_run, size=11)
 
    # ── Sign-off ──────────────────────────────────────────────────────────────
    doc.add_paragraph()   # blank line before sign-off
    signoff_para = doc.add_paragraph()
    signoff_para.paragraph_format.space_after = Pt(2)
    style_run(signoff_para.add_run("Respectfully,"), size=11)
 
    doc.add_paragraph()   # space for signature
    doc.add_paragraph()
 
    name_sig = doc.add_paragraph()
    name_sig.paragraph_format.space_after = Pt(2)
    style_run(name_sig.add_run("Adam Haywood"), size=11, bold=True)
 
    footer_line = doc.add_paragraph()
    style_run(
        footer_line.add_run("U.S. Navy Veteran  |  linkedin.com/in/haywooda"),
        size=9, color=(85, 85, 85)
    )
 
    # ── Metadata section (replaces the === header block from the old .txt) ────
    # This gives you the same "filing" info the old .txt had, tucked at the end.
    doc.add_page_break()
 
    meta_label = doc.add_paragraph()
    style_run(meta_label.add_run("— Reference Copy —"), size=9, color=(150, 150, 150))
 
    meta_job = doc.add_paragraph()
    meta_job.paragraph_format.space_after = Pt(2)
    style_run(meta_job.add_run(f"Job: {job_title}"), size=9, color=(120, 120, 120))
 
    meta_date = doc.add_paragraph()
    meta_date.paragraph_format.space_after = Pt(10)
    style_run(
        meta_date.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y %H:%M')}"),
        size=9, color=(120, 120, 120)
    )
 
    jd_label = doc.add_paragraph()
    jd_label.paragraph_format.space_after = Pt(4)
    style_run(jd_label.add_run("Job Description:"), size=10, bold=True, color=(85, 85, 85))
 
    # Split and add the job description the same way we handled the letter body
    for jd_para in [p.strip() for p in job_description.split("\n\n") if p.strip()]:
        jd_para = " ".join(jd_para.splitlines())
        jd_p = doc.add_paragraph()
        jd_p.paragraph_format.space_after = Pt(6)
        style_run(jd_p.add_run(jd_para), size=9, color=(100, 100, 100))
 
    # ── Write the file to disk ────────────────────────────────────────────────
    # Nothing above actually touched the filesystem — doc.save() is the
    # single call that serializes the whole Document object into a .docx file.
    doc.save(filename)
    return filename
# ─────────────────────────────────────────────────────────────────────────────
 

def main():
    print("🤖 Job Search Assistant — Adam Haywood")
    print("=======================================")
    while True:
        job_description = get_job_description()
        if not job_description.strip():
            print("No job description entered. Exiting.")
            break
        job_title = get_job_title()
        cover_letter = generate_cover_letter(job_description)
        print(f"\n✅ COVER LETTER:\n\n{cover_letter}")
        filename = save_cover_letter(cover_letter, job_description, job_title)
        print(f"\n💾 Saved to: {filename}")
        another = input("\nGenerate another? (y/n): ")
        if another.lower() != "y":
            break
    print_token_summary()  # ← show token cost at end


if __name__ == "__main__":
    main()
